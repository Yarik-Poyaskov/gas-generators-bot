import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def convert_md_to_docx(md_file, docx_file):
    if not os.path.exists(md_file):
        print(f"File {md_file} not found!")
        return

    doc = Document()
    
    with open(md_file, "r", encoding="utf-8") as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Headers
        if line.startswith("# "):
            h = doc.add_heading(line[2:], level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("#### "):
            doc.add_heading(line[5:], level=3)
        # Lists
        elif line.startswith("* ") or line.startswith("- "):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif any(line.startswith(f"{i}. ") for i in range(1, 20)):
            # Simple numbered list handling
            content = line.split(". ", 1)[1]
            doc.add_paragraph(content, style='List Number')
        # Bold text
        elif "**" in line:
            p = doc.add_paragraph()
            parts = line.split("**")
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 != 0:
                    run.bold = True
        # Regular text
        else:
            doc.add_paragraph(line)

    doc.save(docx_file)
    print(f"Successfully created {docx_file}")

if __name__ == "__main__":
    convert_md_to_docx("BOT_MAP.md", "BOT_MAP.docx")
